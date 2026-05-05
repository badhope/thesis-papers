import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
md_path = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Manuscript_Chinese.md'
figures_dir = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Job_Architectures_LLM\Figures'
output_path = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Manuscript_Chinese.docx'

# Read markdown file
with open(md_path, 'r', encoding='utf-8') as f:
    md_content = f.read()

doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.element.set(qn('w:eastAsia'), '宋体')
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5

# Title style
title_style = doc.styles['Title']
title_style.font.name = 'Times New Roman'
title_style.font.element.set(qn('w:eastAsia'), '宋体')
title_style.font.size = Pt(16)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(12)

# Heading styles
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.element.set(qn('w:eastAsia'), '宋体')
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
    elif i == 2:
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True
    else:
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

def add_formatted_paragraph(doc, text):
    """Add a paragraph with formatted text including bold and italic"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Split by ** and * to handle bold and italic
    parts = re.split(r'(\*\*|\*|`)', text)
    
    i = 0
    in_bold = False
    in_italic = False
    in_code = False
    
    while i < len(parts):
        part = parts[i]
        
        if part == '**':
            in_bold = not in_bold
            i += 1
        elif part == '*':
            in_italic = not in_italic
            i += 1
        elif part == '`':
            in_code = not in_code
            i += 1
        else:
            if part.strip():
                run = p.add_run(part)
                if in_bold:
                    run.bold = True
                if in_italic:
                    run.italic = True
                if in_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(11)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.element.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
            i += 1
    
    return p

def add_bullet_point(doc, text):
    """Add a bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.5)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.element.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    
    return p

def add_table_from_markdown(doc, table_text):
    """Convert markdown table to Word table"""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    
    if not lines:
        return None
    
    # Parse table
    data = []
    for line in lines:
        if line.startswith('|---') or line.startswith('| -'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        data.append(cells)
    
    if not data:
        return None
    
    # Create table
    num_rows = len(data)
    num_cols = len(data[0])
    
    # Make sure all rows have same number of columns
    for i, row in enumerate(data):
        while len(row) < num_cols:
            row.append('')
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill data
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Header row formatting
            is_header = i == 0 or (len(data) > 2 and i == 1 and '**' in cell_text)
            
            run = p.add_run(cell_text.replace('**', '').replace('*', ''))
            run.bold = is_header
            run.font.name = 'Times New Roman'
            run.font.element.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
    
    return table

# Process the markdown content
lines = md_content.split('\n')
i = 0
in_table = False
table_lines = []

while i < len(lines):
    line = lines[i]
    
    # Handle figure references - insert actual images
    fig_match = re.search(r'`Figures/(Figure\d+_[\w]+\.png)`', line)
    if fig_match:
        fig_filename = fig_match.group(1)
        fig_path = os.path.join(figures_dir, fig_filename)
        
        # Add the text before the figure reference
        text_before = line[:fig_match.start()].strip()
        text_after = line[fig_match.end():].strip()
        
        if text_before:
            add_formatted_paragraph(doc, text_before)
        
        # Insert the actual image
        if os.path.exists(fig_path):
            # Extract figure number for caption (handle both "图X" and "Figure X")
            fig_num_match = re.search(r'图\s*(\d+)', text_before) or re.search(r'Figure\s*(\d+)', text_before)
            if fig_num_match:
                fig_num = fig_num_match.group(1)
                # Add figure caption
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_p.add_run(f'图 {fig_num}')
                caption_run.bold = True
                caption_run.font.name = 'Times New Roman'
                caption_run.font.element.set(qn('w:eastAsia'), '宋体')
                caption_run.font.size = Pt(11)
                caption_p.paragraph_format.space_after = Pt(6)
            
            # Add the image
            doc.add_picture(fig_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(12)
        
        if text_after:
            add_formatted_paragraph(doc, text_after)
        
        i += 1
        continue
    
    # Handle headings
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).replace('**', '')
        
        if level == 1:
            doc.add_heading(text, level=1)
        elif level == 2:
            doc.add_heading(text, level=2)
        elif level == 3:
            doc.add_heading(text, level=3)
        else:
            add_formatted_paragraph(doc, f'**{text}**')
        
        i += 1
        continue
    
    # Handle tables
    if line.strip().startswith('|'):
        if not in_table:
            in_table = True
            table_lines = []
        table_lines.append(line)
        i += 1
        continue
    else:
        if in_table and table_lines:
            # Process the table
            table_text = '\n'.join(table_lines)
            add_table_from_markdown(doc, table_text)
            in_table = False
            table_lines = []
            
            # Add space after table
            doc.add_paragraph()
    
    # Handle bullet points
    bullet_match = re.match(r'^-\s+(.+)$', line)
    if bullet_match:
        add_bullet_point(doc, bullet_match.group(1))
        i += 1
        continue
    
    # Handle empty lines
    if not line.strip():
        i += 1
        continue
    
    # Handle blockquotes
    if line.startswith('>'):
        text = line[1:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text.replace('**', '').replace('*', ''))
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.element.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
        i += 1
        continue
    
    # Handle regular paragraphs
    if line.strip():
        add_formatted_paragraph(doc, line.strip())
    
    i += 1

# Save the document
doc.save(output_path)
print(f'中文版Word文档已保存至: {output_path}')
