from docx import Document

doc_path = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Manuscript_Chinese.docx'
doc = Document(doc_path)

total_paragraphs = len(doc.paragraphs)
total_tables = len(doc.tables)
total_pictures = 0

for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        total_pictures += 1

print("=" * 80)
print("中文版Word文档内容验证")
print("=" * 80)

print(f"\n基本统计:")
print(f"  段落数量: {total_paragraphs}")
print(f"  表格数量: {total_tables}")
print(f"  图片数量: {total_pictures}")

print(f"\n章节标题:")
headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
for h in headings:
    print(f"  [{h.style.name}] {h.text}")

print(f"\n图片标题:")
captions = [p for p in doc.paragraphs if '图' in p.text]
for c in captions[:5]:
    print(f"  {c.text}")

expected_tables = 13
expected_figures = 3

print(f"\n表格验证:")
print(f"  预期表格: {expected_tables}")
print(f"  实际表格: {total_tables}")
print(f"  {'✓ 所有表格已包含' if total_tables >= expected_tables else '✗ 缺少表格'}")

print(f"\n图片验证:")
print(f"  预期图片: {expected_figures}")
print(f"  实际图片: {total_pictures}")
print(f"  {'✓ 所有图片已包含' if total_pictures >= expected_figures else '✗ 缺少图片'}")

print("\n✓ 验证完成!")
