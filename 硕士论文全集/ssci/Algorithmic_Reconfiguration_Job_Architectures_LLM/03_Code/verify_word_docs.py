from docx import Document

def check_word_doc(doc_path, version_name):
    doc = Document(doc_path)
    
    print("=" * 80)
    print(f"{version_name} - Word文档详细检查")
    print("=" * 80)
    
    errors = []
    warnings = []
    
    # 1. Check tables consistency
    print("\n1. 表格检查")
    print("-" * 60)
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        
        # Check if all rows have same number of cells
        inconsistent_rows = []
        for j, row in enumerate(table.rows):
            if len(row.cells) != cols:
                inconsistent_rows.append(j)
        
        if inconsistent_rows:
            errors.append(f"表格{i+1}: 行 {inconsistent_rows} 列数不一致")
            print(f"  ✗ 表格{i+1}: 列数不一致 - 错误")
        else:
            # Check first row (header) content
            header_text = [cell.text for cell in table.rows[0].cells]
            print(f"  ✓ 表格{i+1}: {rows}行 x {cols}列 - 表头: {header_text[:3]}...")
        
        # Check for empty cells
        empty_cells = 0
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():
                    empty_cells += 1
        
        if empty_cells > 0:
            warnings.append(f"表格{i+1}: 有{empty_cells}个空单元格")
            print(f"  ⚠ 表格{i+1}: {empty_cells}个空单元格")
    
    # 2. Check figures
    print("\n2. 图片检查")
    print("-" * 60)
    picture_count = 0
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            picture_count += 1
    
    print(f"  图片数量: {picture_count}")
    if picture_count < 3:
        errors.append(f"缺少图片: 预期3张，实际{picture_count}张")
        print(f"  ✗ 缺少图片!")
    else:
        print(f"  ✓ 图片数量正确")
    
    # 3. Check for figure captions
    print("\n3. 图片标题检查")
    print("-" * 60)
    fig_captions = []
    for p in doc.paragraphs:
        if '图 ' in p.text or 'Figure ' in p.text:
            if len(p.text) < 50:  # Likely a caption
                fig_captions.append(p.text)
    
    print(f"  找到 {len(fig_captions)} 个图片标题:")
    for cap in fig_captions:
        print(f"    {cap}")
    
    if len(fig_captions) < 3:
        warnings.append(f"图片标题不足: 预期3个，实际{len(fig_captions)}个")
    
    # 4. Check headings structure
    print("\n4. 章节结构检查")
    print("-" * 60)
    headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
    
    expected_sections = ['摘要', '引言', '理论背景', '研究方法', '实证结果', '讨论', '结论', '参考文献', '附录']
    expected_sections_en = ['Abstract', 'Introduction', 'Theoretical', 'Methodology', 'Results', 'Discussion', 'Conclusion', 'References', 'Appendix']
    
    heading_texts = [h.text for h in headings]
    
    for exp_cn, exp_en in zip(expected_sections, expected_sections_en):
        found = any(exp_cn in h or exp_en in h for h in heading_texts)
        if found:
            print(f"  ✓ {exp_cn}/{exp_en}")
        else:
            errors.append(f"缺少章节: {exp_cn}")
            print(f"  ✗ {exp_cn} - 缺失!")
    
    # 5. Check for table notes
    print("\n5. 表格注释检查")
    print("-" * 60)
    note_paragraphs = [p for p in doc.paragraphs if '注：' in p.text or 'Note:' in p.text or '*注' in p.text or '*p <' in p.text]
    print(f"  找到 {len(note_paragraphs)} 个表格注释段落")
    if len(note_paragraphs) < 13:
        warnings.append(f"表格注释可能不足: 预期至少13个，实际{len(note_paragraphs)}个")
    
    # 6. Check for data consistency
    print("\n6. 数据一致性检查")
    print("-" * 60)
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # Check key coefficients
    key_values = {
        '-0.0265': 'H1a系数',
        '-0.0376': 'H1b系数',
        '-0.0058': 'H2系数',
        '20,609': '样本量',
        '7,153': '企业数量',
    }
    
    for value, desc in key_values.items():
        if value in full_text:
            print(f"  ✓ {desc}: {value}")
        else:
            errors.append(f"缺少关键数据: {desc} ({value})")
            print(f"  ✗ {desc}: {value} - 缺失!")
    
    # Final summary
    print("\n" + "=" * 80)
    print(f"检查汇总 - {version_name}")
    print("=" * 80)
    print(f"\n错误: {len(errors)}")
    for i, e in enumerate(errors, 1):
        print(f"  {i}. {e}")
    
    print(f"\n警告: {len(warnings)}")
    for i, w in enumerate(warnings, 1):
        print(f"  {i}. {w}")
    
    if not errors and not warnings:
        print("\n✓ 所有检查通过!")
    elif not errors:
        print(f"\n无严重错误，{len(warnings)}个警告")
    else:
        print(f"\n{len(errors)}个错误需要修复")
    
    return errors, warnings

# Check both versions
en_path = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Manuscript.docx'
cn_path = r'c:\Users\Administrator\Desktop\thesis-papers\硕士论文全集\ssci\Algorithmic_Reconfiguration_Manuscript_Chinese.docx'

en_errors, en_warnings = check_word_doc(en_path, "英文版")
print("\n\n")
cn_errors, cn_warnings = check_word_doc(cn_path, "中文版")
